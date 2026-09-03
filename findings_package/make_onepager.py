"""One-page Word summary of the wobble evaluation."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"c:\Users\HP\Music\AI Projects\wobble_evaluation\findings_package"
CH = os.path.join(ROOT, "charts")
OUT = os.path.join(ROOT, "Wobble_Evaluation_One_Pager.docx")

TEAL = RGBColor(0x0F, 0x52, 0x75)
RED = RGBColor(0xC8, 0x10, 0x2E)
GREY = RGBColor(0x5A, 0x60, 0x66)
INK = RGBColor(0x1A, 0x1D, 0x21)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.5)
sec.bottom_margin = Inches(0.45)
sec.left_margin = Inches(0.65)
sec.right_margin = Inches(0.65)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)
style.font.color.rgb = INK
style.paragraph_format.space_after = Pt(5)
style.paragraph_format.line_spacing = 1.06


def para(text, size=9.5, bold=False, color=INK, italic=False, after=5,
         align=None, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def rule(color="0F5275", size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    pPr = p._p.get_or_add_pPr()
    bd = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:color"), color)
    bot.set(qn("w:space"), "1")
    bd.append(bot)
    pPr.append(bd)
    return p


def shade(cell, hexcol):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcol)
    tcPr.append(sh)


# ---------------------------------------------------------------- masthead
para("FICO COACHING FRAMEWORK  ·  AI SCORING RELIABILITY  ·  SEPTEMBER 2026",
     size=7.5, bold=True, color=TEAL, after=2)
para("The wobble is in the rubric, not the AI", size=19, bold=True, color=INK,
     after=3)
rule()

para(
    "We scored the same classroom observations over and over with the same AI, the same "
    "rubric and the same wording, then measured how often the verdict changed. Nothing "
    "varied between runs except the scorer itself, so any difference is the scorer\u2019s.",
    after=4)
para(
    "The scorer proved consistent. Judgements were identical across 91% of re-scorings, "
    "and reliability held at 0.87 on a 0\u20131 scale \u2014 the range normally treated as "
    "dependable. The instability that remains is not spread evenly across the framework. "
    "It is concentrated in a handful of indicators whose wording admits two defensible "
    "readings of the same lesson.",
    after=7)

doc.add_picture(os.path.join(CH, "onepager_1_concentration.png"), width=Inches(6.15))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.paragraphs[-1].paragraph_format.space_after = Pt(7)

# ---------------------------------------------------------------- why table
para("Why those seven fail \u2014 two causes, two different fixes",
     size=11.5, bold=True, color=TEAL, after=4)

rows = [
    ("The standard bundles two requirements",
     "C4, B2, C6, B9",
     "A real lesson meets one and fails the other, and the rubric never says which "
     "decides. C4 asks for \u201cdeliberate strategies\u2026 diverse students included.\u201d "
     "Runs answering yes cited the teacher calling on students by name; runs answering "
     "no cited that those were volunteers while many stayed silent. Both readings are "
     "correct.",
     "Rewrite: split the standard, or state which clause governs."),
    ("The evidence is not in the audio",
     "D4, D5, D6",
     "Sustained focus during silent work, willingness to attempt, whether pupils handle "
     "materials \u2014 none of this is audible. The scorer is inferring from silence.",
     "No rewrite helps. Needs video, or reword around what can be heard."),
]

t = doc.add_table(rows=1, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = ["Cause", "Indicators", "What is happening, and what fixes it"]
widths = [Inches(1.55), Inches(0.95), Inches(4.5)]
for i, (c, h) in enumerate(zip(t.rows[0].cells, hdr)):
    shade(c, "0F5275")
    c.width = widths[i]
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for cause, codes, what, fix in rows:
    cells = t.add_row().cells
    for i, txt in enumerate([cause, codes, None]):
        cells[i].width = widths[i]
        if txt is None:
            continue
        p = cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(txt)
        r.font.size = Pt(8.5)
        r.bold = (i == 1)
        r.font.color.rgb = RED if i == 1 else INK
    p = cells[2].paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(what + " ")
    r.font.size = Pt(8.5)
    r2 = p.add_run(fix)
    r2.font.size = Pt(8.5)
    r2.bold = True
    r2.font.color.rgb = TEAL

doc.add_paragraph().paragraph_format.space_after = Pt(3)

doc.add_picture(os.path.join(CH, "onepager_2_effect.png"), width=Inches(5.7))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

# ---------------------------------------------------------------- actions
para("What to do", size=11.5, bold=True, color=TEAL, after=3)
for n, txt in enumerate([
    "Rewrite the four compound standards \u2014 C4, B2, C6, B9. Highest-leverage change "
    "available, and it costs nothing but editing time.",
    "Decide about D4, D5 and D6. Either accept they need video, or reword them around "
    "what an audio transcript can actually evidence.",
    "Report the overall lesson score from a single AI pass. Do not report Section B or "
    "Section D from one pass \u2014 their High/Medium/Low band moves between re-scorings "
    "on about half of lessons.",
], 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    r = p.add_run(f"{n}.  ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RED
    r2 = p.add_run(txt)
    r2.font.size = Pt(9)

rule(color="C9CDD2", size=6)
para(
    "Two limits. This measures whether the scorer agrees with itself, not whether it is "
    "right \u2014 an indicator returning the same wrong answer ten times scores perfectly "
    "here. And the failures cannot be spotted by reading the rubric: across all 37 "
    "indicators no textual feature correlated with the disagreement actually observed "
    "(all |\u03c1| < 0.16). Which indicators fail has to be measured, not inspected.",
    size=8, italic=True, color=GREY, after=0)

doc.save(OUT)
print("wrote", OUT)
